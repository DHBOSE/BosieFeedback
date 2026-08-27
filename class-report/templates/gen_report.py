# -*- coding: utf-8 -*-
"""课堂报告 · Word 生成模板（封面 + 课堂内容 + 课堂表现 + 课程纪要 + 四大板块 + 留空脑图）

用法（三步）：
1. 填充「待填充内容」区域 META（封面五要素：标题/学科/授课形式/主讲内容/报告日期）；
2. 在「正文内容」区域用 heading() / para() / bullet() 写（标题带 emoji）：
   一、课堂内容（精简概览）
   二、课堂表现（本课优点 + 待改进）          ← 新增板块
   三、课程纪要（尽量保留原文）  四、课程收获  五、课堂作业
   六、易错分析  七、学习建议  八、知识点脑图
3. 运行：python gen_report.py，得到 report.docx（脑图板块只预留空白位置，无需任何图片）。

硬性规则：封面不写“学生升学目标”等个性化字段；全文不写任何“AI 生成/AI 辅助生成”水印或描述；
封面信息行用无框线表格排版；标题与合适处加 emoji；脑图只预留空白位置（不生成图片）；课堂作业单列。
依赖：python-docx。中文需显式设置 eastAsia=Microsoft YaHei（下方 set_run 已处理）。
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

THEME = (33, 104, 115)      # 深青 #216873
THEME_HEX = "216873"
ACCENT = (230, 90, 90)      # 暖红 #E65A5A
DARK = (51, 51, 51)
GRAY = (120, 120, 120)
FONT = "Microsoft YaHei"


def set_run(run, size=11, bold=False, color=DARK, name=FONT):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def set_style_font(style, name=FONT, size=11, bold=False, color=DARK):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(*color)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def add_hrule(paragraph, color=THEME_HEX, size=10, space=1):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def para(doc, text="", size=11, bold=False, color=DARK, align=None,
         before=0, after=6, style=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if text:
        set_run(p.add_run(text), size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(style="Heading %d" % level)
    pf = p.paragraph_format
    if level == 1:
        set_run(p.add_run(text), size=16, bold=True, color=THEME)
        pf.space_before = Pt(16)
        pf.space_after = Pt(8)
        add_hrule(p)
    elif level == 2:
        set_run(p.add_run(text), size=13, bold=True, color=THEME)
        pf.space_before = Pt(10)
        pf.space_after = Pt(6)
    return p


def bullet(doc, lead, rest, lead_color=ACCENT):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    if lead:
        set_run(p.add_run(lead), size=11, bold=True, color=lead_color)
    if rest:
        set_run(p.add_run(rest), size=11, color=DARK)
    return p


def rich(doc, parts, before=0, after=6, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    for (t, s, b, c) in parts:
        set_run(p.add_run(t), size=s, bold=b, color=c)
    return p


def _run_width(text, size):
    w = 0.0
    for ch in text:
        w += size if ord(ch) > 0x2E80 else size * 0.55
    return w


def _fit_size(text, cell_cm, start=11.5, floor=8.5):
    usable = cell_cm - 0.4
    size = start
    while size > floor and _run_width(text, size) * 0.0353 > usable:
        size -= 0.5
    return size


def cover_table(doc, info):
    """封面信息行：无框线两列表格（标签右对齐、值左对齐，整表居中）。"""
    table = doc.add_table(rows=len(info), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        borders.append(el)
    table._tbl.tblPr.append(borders)
    for i, (label, value) in enumerate(info):
        cl = table.cell(i, 0)
        cv = table.cell(i, 1)
        cl.width = Cm(4.0)
        cv.width = Cm(8.2)
        pl = cl.paragraphs[0]
        pl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pl.paragraph_format.space_after = Pt(8)
        set_run(pl.add_run(label), size=11.5, bold=True, color=THEME)
        pv = cv.paragraphs[0]
        pv.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pv.paragraph_format.space_after = Pt(8)
        set_run(pv.add_run(value), size=_fit_size(value, 8.2), bold=False, color=DARK)
    return table


# ================== 待填充内容：封面信息（固定五要素） ==================
META = {
    "title": "课程标题",                            # 主标题
    "subtitle": "—— 章节副标题 / 说明",              # 副标题
    "info": [                                        # 信息行（成对：标签, 值），仅这四项
        ("学科", "XX学科（年级 · 章节）"),
        ("授课形式", "一对一辅导课"),
        ("主讲内容", "……"),
        ("报告日期", "20XX 年 X 月 X 日"),
    ],
}
# ====================================================================


doc = Document()
sec = doc.sections[0]
sec.page_width = Mm(210)
sec.page_height = Mm(297)
sec.top_margin = Cm(2.2)
sec.bottom_margin = Cm(2.2)
sec.left_margin = Cm(2.2)
sec.right_margin = Cm(2.2)

set_style_font(doc.styles["Normal"], size=11, color=DARK)
for sn in ["Heading 1", "Heading 2", "Heading 3"]:
    set_style_font(doc.styles[sn], size=16, bold=True, color=THEME)
for sn in ["List Bullet", "List Number"]:
    set_style_font(doc.styles[sn], size=11, color=DARK)

# ---------- 封面（无 AI 水印，信息行用无框线表格） ----------
for _ in range(7):
    para(doc, "", after=0)
para(doc, "🎓 课堂报告", size=13, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
para(doc, META["title"], size=40, bold=True, color=THEME,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
para(doc, META["subtitle"], size=15, color=(90, 90, 90),
     align=WD_ALIGN_PARAGRAPH.CENTER, after=22)
div = para(doc, "", after=22)
add_hrule(div, color="E65A5A", size=14, space=2)
cover_table(doc, META["info"])

pb = doc.add_paragraph()
pb.add_run().add_break(WD_BREAK.PAGE)

# ================== 待填充内容：正文（固定结构，标题带 emoji） ==================
# 一、课堂内容（精简概览）：只写主线 + 4~6 条核心结论 bullet
heading(doc, "📚 一、课堂内容", 1)
para(doc, "本节围绕「……主线……」展开，核心结论如下：")
bullet(doc, "结论标签：", "……")

# 二、课堂表现：本课优点 + 待改进
heading(doc, "🌟 二、课堂表现", 1)
para(doc, "本节课整体表现反馈如下：", after=6)
heading(doc, "👍 优点", 2)
bullet(doc, "标签：", "……")
heading(doc, "🔧 待改进", 2)
bullet(doc, "标签：", "……")

# 三、课程纪要：大幅保留原始纪要，按知识点分小节（小节标题也可加 emoji）
heading(doc, "📝 三、课程纪要", 1)
heading(doc, "🔭 1. 知识点小节标题", 2)
bullet(doc, "", "……保留原始纪要细节、关键数据、实验、推导……")

# 四、课程收获
heading(doc, "🏆 四、课程收获", 1)
bullet(doc, "标签：", "……")

# 五、课堂作业：从纪要「作业布置」提取，含作业范围 + 答题规范提醒
heading(doc, "📋 五、课堂作业", 1)
bullet(doc, "作业：", "……")
bullet(doc, "规范：", "……")

# 六、易错分析：每条「① 误区 + 正确理解」
heading(doc, "⚠️ 六、易错分析", 1)
bullet(doc, "① 误区标签：", "正确理解……")

# 七、学习建议：只写学科学习方法
heading(doc, "💡 七、学习建议", 1)
bullet(doc, "建议：", "……")

# 八、知识点脑图：不生成脑图图片，只预留空白位置，由用户自行添加
heading(doc, "🧠 八、知识点脑图", 1)
para(doc, "（此处为知识点脑图预留空白位置，请自行添加脑图。）", size=10.5, color=GRAY,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
for _ in range(14):
    para(doc, "", after=16)
# ================== 正文内容结束 ==================

doc.save("report.docx")
print("saved: report.docx")