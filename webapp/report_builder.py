# -*- coding: utf-8 -*-
"""课堂报告 Word 生成器（函数版，源自 class-report/templates/gen_report.py）

build_report(data, mindmap_path, out_path)
  data: DeepSeek 返回的结构化 JSON（dict），字段见 PROMPT 中的 JSON Schema
  mindmap_path: 脑图图片路径（可为 None，此时按 SKILL 规范留空白占位）
"""
import os
from docx import Document
from docx.image.image import Image as DocxImage
from docx.shared import Pt, RGBColor, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

THEME = (33, 104, 115)      # 深青 #216873
THEME_HEX = "216873"
ACCENT = (230, 90, 90)      # 暖红 #E65A5A
DARK = (51, 51, 51)
GRAY = (120, 120, 120)
FONT = "Microsoft YaHei"


def set_run(run, size=11, bold=False, color=DARK, name=FONT):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def set_style_font(style, name=FONT, size=11, bold=False, color=DARK):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(*color)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def add_hrule(paragraph, color=THEME_HEX, size=10, space=1):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def para(doc, text="", size=11, bold=False, color=DARK, align=None,
         before=0, after=6, style=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if text:
        set_run(p.add_run(text), size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(style="Heading %d" % level)
    pf = p.paragraph_format
    if level == 1:
        set_run(p.add_run(text), size=16, bold=True, color=THEME)
        pf.space_before = Pt(16)
        pf.space_after = Pt(8)
        add_hrule(p)
    elif level == 2:
        set_run(p.add_run(text), size=13, bold=True, color=THEME)
        pf.space_before = Pt(10)
        pf.space_after = Pt(6)
    return p


def bullet(doc, lead, rest, lead_color=ACCENT):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    if lead:
        set_run(p.add_run(lead), size=11, bold=True, color=lead_color)
    if rest:
        set_run(p.add_run(rest), size=11, color=DARK)
    return p


def _run_width(text, size):
    w = 0.0
    for ch in text:
        w += size if ord(ch) > 0x2E80 else size * 0.55
    return w


def _fit_size(text, cell_cm, start=11.5, floor=8.5):
    usable = cell_cm - 0.4
    size = start
    while size > floor and _run_width(text, size) * 0.0353 > usable:
        size -= 0.5
    return size


MINDMAP_MAX_W_CM = 15.5   # 版心宽度上限
MINDMAP_MAX_H_CM = 21.5   # 高度上限（等比缩放）
CONTENT_W_CM = 16.6       # 版心宽度 = 21 - 2.2*2


def _mindmap_size(path):
    """等比缩放：默认顶满宽度 15.5cm，若高度超过 21.5cm 则改为顶满高度。"""
    img = DocxImage.from_file(path)
    ratio = img.px_width / img.px_height  # 宽 / 高
    w = MINDMAP_MAX_W_CM
    h = w / ratio
    if h > MINDMAP_MAX_H_CM:
        h = MINDMAP_MAX_H_CM
        w = h * ratio
    return w, h


def _upgrade_picture_to_svg(doc, run, svg_path, index=0):
    """把刚 add_picture 嵌入的 PNG 图片升级为 SVG 矢量（PNG 保留为兜底）。

    Word/WPS 支持 docx 内 SVG：主 blip 仍是 PNG（旧阅读器、docx-preview 用它），
    矢量数据放在 a:extLst 的 asvg:svgBlip 扩展里。Word 转 PDF 时按矢量渲染，
    任意缩放都清晰。失败静默降级为纯 PNG。
    """
    try:
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.oxml import parse_xml

        with open(svg_path, "rb") as f:
            svg_bytes = f.read()
        blips = run._r.xpath(".//a:blip")
        if not blips:
            return
        blip = blips[0]
        partname = PackURI("/word/media/mindmap-vector-%d.svg" % index)
        svg_part = Part(partname, "image/svg+xml", svg_bytes, doc.part.package)
        rId = doc.part.relate_to(svg_part, RT.IMAGE)
        ext_xml = (
            '<a:extLst xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
            'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
            'xmlns:asvg="http://schemas.microsoft.com/office/drawing/2016/SVG/main">'
            '<a:ext uri="{96DAC541-7B7A-43D3-8B79-37D633B846F1}">'
            '<asvg:svgBlip r:embed="%s"/>'
            "</a:ext></a:extLst>" % rId)
        blip.append(parse_xml(ext_xml))
    except Exception:
        pass


# ---------------- 机构品牌 / 水印 ----------------
SLOT_DEFAULT_H = {"cover": 1.2, "header": 0.7, "footer": 0.5}


def _strip_borders(table):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        borders.append(el)
    table._tbl.tblPr.append(borders)


def _slot(brand, name):
    """取某槽位（cover/header/footer）配置：enabled / height_cm / path。"""
    s = ((brand or {}).get(name)) or {}
    try:
        h = float(s.get("height_cm") or SLOT_DEFAULT_H[name])
    except (TypeError, ValueError):
        h = SLOT_DEFAULT_H[name]
    return {"enabled": bool(s.get("enabled", True)),
            "height_cm": max(0.3, min(8.0, h)),
            "path": s.get("path")}


def _brand_active(brand):
    if not brand:
        return False
    if brand.get("org_name") or brand.get("contact"):
        return True
    return any(_slot(brand, s)["path"] for s in ("cover", "header", "footer"))


def cover_brand(doc, brand, slot):
    """封面顶部：图片 + 机构名横排（整体居中），图片高度按 slot 设置。"""
    logo = slot["path"]
    h = Cm(slot["height_cm"])
    name = brand.get("org_name") or ""
    if logo and name:
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        _strip_borders(table)
        cl, cv = table.cell(0, 0), table.cell(0, 1)
        cl.width = Cm(slot["height_cm"] + 0.8)
        cv.width = Cm(8.0)
        for c in (cl, cv):
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pl = cl.paragraphs[0]
        pl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pl.paragraph_format.space_after = Pt(0)
        pl.add_run().add_picture(logo, height=h)
        pv = cv.paragraphs[0]
        pv.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pv.paragraph_format.space_after = Pt(0)
        set_run(pv.add_run("  " + name), size=15, bold=True, color=THEME)
    elif logo:
        p = para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
        p.add_run().add_picture(logo, height=h)
    else:
        para(doc, name, size=15, bold=True, color=THEME,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=0)


def _page_number_run(paragraph):
    """插入 PAGE 域（页码）。"""
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")  # 9pt
    rPr.append(sz)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "787878")
    rPr.append(color)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


def _top_border(paragraph, color="D9D9D9", size=6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), str(size))
    top.set(qn("w:space"), "4")
    top.set(qn("w:color"), color)
    pBdr.append(top)
    pPr.append(pBdr)


def _restart_page_numbering(sec):
    """该节页码从 1 重新开始（封面不计页码）。"""
    sectPr = sec._sectPr
    pg = sectPr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        cols = sectPr.find(qn("w:cols"))
        if cols is not None:
            cols.addprevious(pg)
        else:
            sectPr.append(pg)
    pg.set(qn("w:start"), "1")


def setup_header_footer(sec, brand, header_slot, footer_slot):
    """页眉/页脚内容全部靠右：
    页眉：机构名 + 图片（下细分隔线）；
    页脚：机构名 · 联系方式 · 第 X 页 · 图片（上细分隔线）。
    作用于正文节（封面节无页眉页脚，页码从正文起算第 1 页）。"""
    name = brand.get("org_name") or ""
    contact = brand.get("contact") or ""

    # ---- 页眉 ----
    if header_slot["enabled"] and (header_slot["path"] or name):
        sec.header.is_linked_to_previous = False
        hp = sec.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.paragraph_format.space_after = Pt(2)
        if name:
            set_run(hp.add_run(name + ("  " if header_slot["path"] else "")),
                    size=10.5, bold=True, color=THEME)
        if header_slot["path"]:
            hp.add_run().add_picture(header_slot["path"], height=Cm(header_slot["height_cm"]))
        add_hrule(hp, color="D9D9D9", size=6)

    # ---- 页脚 ----
    if footer_slot["enabled"]:
        sec.footer.is_linked_to_previous = False
        fp = sec.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp.paragraph_format.space_before = Pt(2)
        parts = []
        if name:
            parts.append(name)
        if contact:
            parts.append(contact)
        if parts:
            set_run(fp.add_run("  ·  ".join(parts) + "  ·  "), size=9, color=GRAY)
        set_run(fp.add_run("第 "), size=9, color=GRAY)
        _page_number_run(fp)
        set_run(fp.add_run(" 页"), size=9, color=GRAY)
        if footer_slot["path"]:
            fp.add_run().add_text("  ")
            fp.add_run().add_picture(footer_slot["path"], height=Cm(footer_slot["height_cm"]))
        _top_border(fp)


# ---------------- 正文大水印（浮动图片，文字下方居中） ----------------
WM_DEFAULTS = {"enabled": False, "type": "text", "text": "",
               "size_pct": 60, "opacity_pct": 15, "angle": -45,
               "font": "msyh", "color": THEME_HEX, "path": None}
# 水印文字可选字体：(key, 显示名, 字体文件名, ttc 内 index, CSS font-family)
# 覆盖 Windows 自带 + Office 华文/方正系列；按文件是否存在过滤
WM_FONTS_DIR = r"C:\Windows\Fonts"
WM_FONT_TABLE = [
    ("msyh",      "微软雅黑",      "msyh.ttc",     0, "Microsoft YaHei"),
    ("msyhl",     "微软雅黑细体",  "msyhl.ttc",    0, "Microsoft YaHei Light"),
    ("simhei",    "黑体",          "simhei.ttf",   0, "SimHei"),
    ("simsun",    "宋体",          "simsun.ttc",   0, "SimSun"),
    ("nsimsun",   "新宋体",        "simsun.ttc",   1, "NSimSun"),
    ("simfang",   "仿宋",          "simfang.ttf",  0, "FangSong"),
    ("simkai",    "楷体",          "simkai.ttf",   0, "KaiTi"),
    ("deng",      "等线",          "Deng.ttf",     0, "DengXian"),
    ("dengl",     "等线细体",      "Dengl.ttf",    0, "DengXian Light"),
    ("lisu",      "隶书",          "SIMLI.TTF",    0, "LiSu"),
    ("youyuan",   "幼圆",          "SIMYOU.TTF",   0, "YouYuan"),
    ("stsong",    "华文宋体",      "STSONG.TTF",   0, "STSong"),
    ("stzhongs",  "华文中宋",      "STZHONGS.TTF", 0, "STZhongsong"),
    ("stkaiti",   "华文楷体",      "STKAITI.TTF",  0, "STKaiti"),
    ("stfangso",  "华文仿宋",      "STFANGSO.TTF", 0, "STFangsong"),
    ("stxihei",   "华文细黑",      "STXIHEI.TTF",  0, "STXihei"),
    ("stliti",    "华文隶书",      "STLITI.TTF",   0, "STLiti"),
    ("stxingkai", "华文行楷",      "STXINGKA.TTF", 0, "STXingkai"),
    ("stxinwei",  "华文新魏",      "STXINWEI.TTF", 0, "STXinwei"),
    ("sthupo",    "华文琥珀",      "STHUPO.TTF",   0, "STHupo"),
    ("stcaiyun",  "华文彩云",      "STCAIYUN.TTF", 0, "STCaiyun"),
    ("fzstk",     "方正舒体",      "FZSTK.TTF",    0, "FZShuTi"),
    ("fzytk",     "方正姚体",      "FZYTK.TTF",    0, "FZYaoTi"),
]
WM_FONT_KEYS = set(k for k, _, _, _, _ in WM_FONT_TABLE)


def available_wm_fonts():
    """本机实际存在的水印字体清单（前端下拉用）。"""
    return [{"key": k, "label": label, "css": css}
            for k, label, fn, idx, css in WM_FONT_TABLE
            if os.path.isfile(os.path.join(WM_FONTS_DIR, fn))]


def _wm_font_file(key):
    """按 key 解析字体文件 (path, ttc_index)；不存在则回退到第一个可用字体。"""
    for k, _, fn, idx, _ in WM_FONT_TABLE:
        if k == key:
            p = os.path.join(WM_FONTS_DIR, fn)
            if os.path.isfile(p):
                return p, idx
            break
    for _, _, fn, idx, _ in WM_FONT_TABLE:
        p = os.path.join(WM_FONTS_DIR, fn)
        if os.path.isfile(p):
            return p, idx
    return None


def _watermark_cfg(brand):
    wm = dict(WM_DEFAULTS)
    wm.update(((brand or {}).get("watermark")) or {})
    if wm.get("type") not in ("text", "image"):
        wm["type"] = "text"
    def _num(key, lo, hi, default):
        try:
            v = float(wm.get(key))
        except (TypeError, ValueError):
            v = default
        return max(lo, min(hi, v))
    wm["size_pct"] = _num("size_pct", 10, 200, 60)
    wm["opacity_pct"] = _num("opacity_pct", 3, 100, 15)
    wm["angle"] = _num("angle", -180, 180, -45)
    wm["enabled"] = bool(wm.get("enabled"))
    wm["text"] = str(wm.get("text") or "").strip()
    wm["font"] = str(wm.get("font") or "msyh")
    if wm["font"] not in WM_FONT_KEYS:
        wm["font"] = "msyh"
    c = str(wm.get("color") or "").lstrip("#").lower()
    if len(c) != 6 or any(ch not in "0123456789abcdef" for ch in c):
        c = THEME_HEX.lower()
    wm["color"] = c
    return wm


def _render_watermark_image(wm, out_path):
    """把水印（文字/图片 + 角度 + 透明度）烘焙成一张透明底 PNG。
    返回 (path, w_px, h_px)；无法生成返回 None。PIL 缺失/字体缺失都不阻断报告生成。"""
    try:
        from PIL import Image, ImageDraw, ImageFont
    except Exception:
        return None
    try:
        if wm["type"] == "image":
            if not wm.get("path") or not os.path.isfile(wm["path"]):
                return None
            img = Image.open(wm["path"]).convert("RGBA")
            if img.width > 2400:  # 像素只影响清晰度，限制体积
                img = img.resize((2400, int(img.height * 2400 / img.width)),
                                 Image.LANCZOS)
        else:
            if not wm["text"]:
                return None
            ff = _wm_font_file(wm.get("font"))
            if not ff:
                return None
            try:
                font = ImageFont.truetype(ff[0], 300, index=ff[1])
            except Exception:
                return None
            probe = ImageDraw.Draw(Image.new("RGBA", (8, 8)))
            bbox = probe.textbbox((0, 0), wm["text"], font=font)
            w, h = bbox[2] - bbox[0], bbox[3] - bbox[1]
            pad = 30
            img = Image.new("RGBA", (w + pad * 2, h + pad * 2), (0, 0, 0, 0))
            rgb = tuple(int(wm["color"][i:i + 2], 16) for i in (0, 2, 4))
            ImageDraw.Draw(img).text(
                (pad - bbox[0], pad - bbox[1]), wm["text"], font=font,
                fill=rgb + (255,))
        if wm["angle"]:
            img = img.rotate(wm["angle"], expand=True, resample=Image.BICUBIC)
        op = wm["opacity_pct"] / 100.0
        img.putalpha(img.getchannel("A").point(lambda a: int(a * op)))
        img.save(out_path)
        return out_path, img.width, img.height
    except Exception:
        return None


def _float_behind_run(run, image_path, width_cm, height_cm, docpr_id, name):
    """在 run 中插入浮动图片：behindDoc（文字下方），相对页面水平/垂直居中。"""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    rId, _ = run.part.get_or_add_image(image_path)
    cx, cy = int(Cm(width_cm)), int(Cm(height_cm))
    xml = (
        '<w:drawing %s>'
        '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" '
        'relativeHeight="0" behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">'
        '<wp:simplePos x="0" y="0"/>'
        '<wp:positionH relativeFrom="page"><wp:align>center</wp:align></wp:positionH>'
        '<wp:positionV relativeFrom="page"><wp:align>center</wp:align></wp:positionV>'
        '<wp:extent cx="%d" cy="%d"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        '<wp:wrapNone/>'
        '<wp:docPr id="%d" name="%s"/>'
        '<wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect="1"/></wp:cNvGraphicFramePr>'
        '<a:graphic><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<pic:pic>'
        '<pic:nvPicPr><pic:cNvPr id="%d" name="%s"/>'
        '<pic:cNvPicPr><a:picLocks noChangeAspect="1" noChangeArrowheads="1"/></pic:cNvPicPr></pic:nvPicPr>'
        '<pic:blipFill><a:blip r:embed="%s"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill>'
        '<pic:spPr bwMode="auto"><a:xfrm><a:off x="0" y="0"/><a:ext cx="%d" cy="%d"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr>'
        '</pic:pic>'
        '</a:graphicData></a:graphic>'
        '</wp:anchor>'
        '</w:drawing>'
    ) % (nsdecls("w", "wp", "a", "pic", "r"), cx, cy, docpr_id, name,
         docpr_id, name, rId, cx, cy)
    run._r.append(parse_xml(xml))


def _add_body_watermark(sec, brand):
    """正文大水印：作用于该节每一页（含页眉未启用时），封面节不受影响。"""
    wm = _watermark_cfg(brand)
    if not wm["enabled"]:
        return
    if wm["type"] == "text" and not wm["text"]:
        wm["text"] = ((brand or {}).get("org_name") or "").strip()
    import tempfile
    fd, tmp = tempfile.mkstemp(suffix=".png", prefix="kf_body_wm_")
    os.close(fd)
    try:
        rendered = _render_watermark_image(wm, tmp)
        if not rendered:
            return
        _, w_px, h_px = rendered
        content_w_cm = (sec.page_width - sec.left_margin
                        - sec.right_margin) / 360000.0
        target_w_cm = content_w_cm * wm["size_pct"] / 100.0
        target_h_cm = target_w_cm * h_px / float(w_px)
        hdr = sec.header
        hdr.is_linked_to_previous = False
        p = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        _float_behind_run(p.add_run(), tmp, target_w_cm, target_h_cm,
                          88000001, "正文水印")
    finally:
        if os.path.isfile(tmp):
            os.remove(tmp)


def cover_table(doc, info):
    """封面信息行：无框线两列表格（标签右对齐、值左对齐，整表居中）。"""
    table = doc.add_table(rows=len(info), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        borders.append(el)
    table._tbl.tblPr.append(borders)
    for i, (label, value) in enumerate(info):
        cl = table.cell(i, 0)
        cv = table.cell(i, 1)
        cl.width = Cm(4.0)
        cv.width = Cm(8.2)
        pl = cl.paragraphs[0]
        pl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pl.paragraph_format.space_after = Pt(8)
        set_run(pl.add_run(label), size=11.5, bold=True, color=THEME)
        pv = cv.paragraphs[0]
        pv.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pv.paragraph_format.space_after = Pt(8)
        set_run(pv.add_run(value), size=_fit_size(value, 8.2), bold=False, color=DARK)
    return table


def build_report(data, mindmap_path, out_path, brand=None):
    meta = data["meta"]
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)

    set_style_font(doc.styles["Normal"], size=11, color=DARK)
    for sn in ["Heading 1", "Heading 2", "Heading 3"]:
        set_style_font(doc.styles[sn], size=16, bold=True, color=THEME)
    for sn in ["List Bullet", "List Number"]:
        set_style_font(doc.styles[sn], size=11, color=DARK)

    # ---------- 封面（第一节：无页眉页脚、不计页码） ----------
    cover_slot = _slot(brand, "cover")
    if _brand_active(brand) and cover_slot["enabled"] and \
            (cover_slot["path"] or brand.get("org_name")):
        cover_brand(doc, brand, cover_slot)
        for _ in range(5):
            para(doc, "", after=0)
    else:
        for _ in range(7):
            para(doc, "", after=0)
    para(doc, "🎓 课堂报告", size=13, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
    para(doc, meta["title"], size=40, bold=True, color=THEME,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    para(doc, meta["subtitle"], size=15, color=(90, 90, 90),
         align=WD_ALIGN_PARAGRAPH.CENTER, after=22)
    div = para(doc, "", after=22)
    add_hrule(div, color="E65A5A", size=14, space=2)
    info = [("学科", meta["subject"])]
    if meta.get("student"):
        info.append(("学生", meta["student"]))
    info += [("授课形式", meta["form"]),
             ("主讲内容", meta["content"]),
             ("报告日期", meta["date"])]
    cover_table(doc, info)

    # ---------- 正文节：新起一节，页码从 1 开始，页眉页脚仅作用于本节 ----------
    sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
    _restart_page_numbering(sec2)
    if _brand_active(brand):
        setup_header_footer(sec2, brand, _slot(brand, "header"), _slot(brand, "footer"))
    _add_body_watermark(sec2, brand)

    # 一、课堂内容
    heading(doc, "📚 一、课堂内容", 1)
    para(doc, data.get("mainline", ""))
    for it in data["content"]:
        bullet(doc, it.get("lead", ""), it.get("text", ""))

    # 二、课堂表现
    heading(doc, "🌟 二、课堂表现", 1)
    para(doc, "本节课整体表现反馈如下：", after=6)
    heading(doc, "👍 优点", 2)
    for it in data["performance"]["pros"]:
        bullet(doc, it.get("lead", ""), it.get("text", ""))
    heading(doc, "🔧 待改进", 2)
    for it in data["performance"]["cons"]:
        bullet(doc, it.get("lead", ""), it.get("text", ""))

    # 三、课程纪要
    heading(doc, "📝 三、课程纪要", 1)
    for sub in data["minutes"]:
        heading(doc, sub["title"], 2)
        for pt in sub["points"]:
            bullet(doc, "", pt)

    # 四、课程收获
    heading(doc, "🏆 四、课程收获", 1)
    for it in data["gains"]:
        bullet(doc, it.get("lead", ""), it.get("text", ""))

    # 五、课堂作业
    heading(doc, "📋 五、课堂作业", 1)
    for it in data["homework"]:
        bullet(doc, it.get("lead", ""), it.get("text", ""))

    # 六、易错分析
    heading(doc, "⚠️ 六、易错分析", 1)
    for i, it in enumerate(data["mistakes"], 1):
        bullet(doc, "%d %s" % (i, it.get("lead", "")), it.get("text", ""))

    # 七、学习建议
    heading(doc, "💡 七、学习建议", 1)
    for it in data["suggestions"]:
        bullet(doc, it.get("lead", ""), it.get("text", ""))

    # 八、知识点脑图：单图直接嵌入；AI 生成的切片图逐页嵌入（每段顶满版心宽度）
    heading(doc, "🧠 八、知识点脑图", 1)
    paths = mindmap_path if isinstance(mindmap_path, list) else ([mindmap_path] if mindmap_path else [])
    paths = [p for p in paths if p and os.path.isfile(p)]
    if paths:
        if len(paths) > 1:
            para(doc, "（脑图较大，已按知识分支分段排版，连续 %d 页，顺序阅读。）" % len(paths),
                 size=10.5, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
        for i, mp in enumerate(paths):
            if i > 0:
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            w, h = _mindmap_size(mp)
            run.add_picture(mp, width=Cm(w), height=Cm(h))
            # 有配套 SVG（AI 生成脑图切片）则升级为矢量嵌入，PDF 任意缩放不糊
            svg_path = os.path.splitext(mp)[0] + ".svg"
            if os.path.isfile(svg_path):
                _upgrade_picture_to_svg(doc, run, svg_path, index=i)
    else:
        para(doc, "（此处为知识点脑图预留空白位置，请自行添加脑图。）", size=10.5,
             color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
        for _ in range(14):
            para(doc, "", after=16)

    doc.save(out_path)
    return out_path


GREEN = (46, 125, 50)   # 优点绿 #2E7D32


def build_student_report(name, reports, summary, out_path, brand=None):
    """学情趋势 Word：封面（学生名 + 统计）→ AI 学情总结（如有）→ 误区汇总 → 逐课时间线。

    name: 学生姓名；reports: [data, ...]（按时间正序）；summary: AI 学情总结纯文本（可空）。
    版式与 build_report 一致（封面节无页眉页脚，正文节页码从 1 起）。
    """
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)

    set_style_font(doc.styles["Normal"], size=11, color=DARK)
    for sn in ["Heading 1", "Heading 2", "Heading 3"]:
        set_style_font(doc.styles[sn], size=16, bold=True, color=THEME)
    for sn in ["List Bullet", "List Number"]:
        set_style_font(doc.styles[sn], size=11, color=DARK)

    # ---------- 封面（第一节：无页眉页脚、不计页码） ----------
    cover_slot = _slot(brand, "cover")
    if _brand_active(brand) and cover_slot["enabled"] and \
            (cover_slot["path"] or brand.get("org_name")):
        cover_brand(doc, brand, cover_slot)
        for _ in range(5):
            para(doc, "", after=0)
    else:
        for _ in range(7):
            para(doc, "", after=0)
    para(doc, "🎓 学情趋势报告", size=13, color=GRAY,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
    para(doc, name, size=40, bold=True, color=THEME,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    first_meta = reports[0].get("meta") or {}
    last_meta = reports[-1].get("meta") or {}
    date_range = first_meta.get("date", "—")
    if len(reports) > 1:
        date_range = "%s ~ %s" % (first_meta.get("date", "—"),
                                  last_meta.get("date", "—"))
    subjects = []
    for r in reports:
        s = (r.get("meta") or {}).get("subject", "")
        if s and s not in subjects:
            subjects.append(s)
    para(doc, "跨 %d 节课的阶段性学情汇总" % len(reports), size=15,
         color=(90, 90, 90), align=WD_ALIGN_PARAGRAPH.CENTER, after=22)
    div = para(doc, "", after=22)
    add_hrule(div, color="E65A5A", size=14, space=2)
    cover_table(doc, [
        ("学生", name),
        ("报告数量", "%d 节课" % len(reports)),
        ("时间范围", date_range),
        ("涉及学科", "、".join(subjects) or "—"),
    ])

    # ---------- 正文节：页码从 1 开始，页眉页脚仅作用于本节 ----------
    sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
    _restart_page_numbering(sec2)
    if _brand_active(brand):
        setup_header_footer(sec2, brand, _slot(brand, "header"),
                            _slot(brand, "footer"))
    _add_body_watermark(sec2, brand)

    # 一、AI 学情总结（可选）
    if summary and summary.strip():
        heading(doc, "🤖 一、AI 学情总结", 1)
        for line in summary.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line[0] in "・-•":
                bullet(doc, "", line.lstrip("・-• ").strip())
            else:
                para(doc, line, size=12, bold=True, color=THEME,
                     before=6, after=4)

    # 二、待改进汇总（跨全部课程，标注日期便于回溯；以课堂表现为依据，不含误区）
    heading(doc, "🔧 二、待改进汇总（全部课程）", 1)
    total = 0
    for r in reports:
        m = r.get("meta") or {}
        for it in (r.get("performance") or {}).get("cons") or []:
            total += 1
            lead = "【%s】%s" % (m.get("date", ""), it.get("lead", ""))
            bullet(doc, lead, it.get("text", ""))
    if not total:
        para(doc, "（各次报告中均未记录待改进项。）", color=GRAY)

    # 三、逐课时间线
    heading(doc, "🗂 三、逐课时间线", 1)
    for r in reports:
        m = r.get("meta") or {}
        heading(doc, "%s · %s" % (m.get("date", ""), m.get("title", "")), 2)
        if m.get("subject"):
            para(doc, m["subject"], size=10, color=GRAY, after=4)
        perf = r.get("performance") or {}
        if perf.get("pros"):
            para(doc, "👍 优点", size=11, bold=True, color=GREEN, after=2)
            for it in perf["pros"]:
                bullet(doc, it.get("lead", ""), it.get("text", ""),
                       lead_color=GREEN)
        if perf.get("cons"):
            para(doc, "🔧 待改进", size=11, bold=True, color=ACCENT, after=2)
            for it in perf["cons"]:
                bullet(doc, it.get("lead", ""), it.get("text", ""))
        if r.get("homework"):
            para(doc, "📋 作业", size=11, bold=True, color=THEME, after=2)
            for it in r["homework"]:
                bullet(doc, it.get("lead", ""), it.get("text", ""),
                       lead_color=THEME)

    doc.save(out_path)
    return out_path
